import httpx
import pdfplumber
from docx import Document
from pathlib import Path

def download_attachment(url: str, cookies: dict, output_path: Path):
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Referer": "https://lms.vit.ac.in/"
    }
    with httpx.Client(cookies=cookies, headers=headers, follow_redirects=True, verify=False, timeout=30.0) as client:
        response = client.get(url)
        response.raise_for_status()
        with open(output_path, "wb") as f:
            f.write(response.content)

def extract_text(pdf_path: Path) -> str:
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
    return text

def scaffold_markdown_to_docx(md_text: str, title: str, output_path: Path):
    doc = Document()
    doc.add_heading(title, level=0)
    
    lines = md_text.split("\n")
    for line in lines:
        cleaned_line = line.replace("###", "").replace("**", "").strip()
        if not cleaned_line:
            continue
            
        if cleaned_line.startswith("- ") or cleaned_line.startswith("* "):
            doc.add_paragraph(cleaned_line[2:], style='List Bullet')
        elif len(cleaned_line) < 50 and not cleaned_line.endswith("."):
            doc.add_heading(cleaned_line, level=2)
        else:
            doc.add_paragraph(cleaned_line)
            
    doc.save(output_path)
